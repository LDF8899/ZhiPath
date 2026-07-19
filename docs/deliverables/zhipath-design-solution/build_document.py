from __future__ import annotations

import math
import os
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SCREENSHOTS = ROOT / "screenshots"
DIAGRAMS = ROOT / "diagrams"
OUTPUT_DOCX = ROOT / "智途ZhiPath_作品设计实现方案_修订完整版.docx"

BLUE = "2563EB"
BLUE_DARK = "1E4F9A"
BLUE_LIGHT = "EAF2FF"
ORANGE = "E64A2E"
GREEN = "2F8F5B"
INK = "27303F"
GRAY = "667085"
LINE = "D0D5DD"

FONT_YAHEI = r"C:\Windows\Fonts\msyh.ttc"
FONT_YAHEI_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"
FONT_SONG = r"C:\Windows\Fonts\simsun.ttc"


def font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size)


def wrap_cn(text: str, max_chars: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        if not raw:
            lines.append("")
            continue
        line = ""
        for ch in raw:
            line += ch
            if len(line) >= max_chars:
                lines.append(line)
                line = ""
        if line:
            lines.append(line)
    return lines


def draw_text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str,
                     fnt: ImageFont.FreeTypeFont, fill: str = "#27303F", max_chars: int = 12,
                     spacing: int = 8) -> None:
    lines = wrap_cn(text, max_chars)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total = sum(heights) + max(0, len(lines) - 1) * spacing
    y = box[1] + (box[3] - box[1] - total) / 2
    for line, h in zip(lines, heights):
        w = draw.textbbox((0, 0), line, font=fnt)[2]
        x = box[0] + (box[2] - box[0] - w) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += h + spacing


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str,
                subtitle: str = "", fill: str = "#FFFFFF", outline: str = "#2563EB",
                title_color: str = "#1E4F9A", width: int = 3) -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=width)
    title_font = font(FONT_YAHEI_BOLD, 30)
    subtitle_font = font(FONT_YAHEI, 20)
    if subtitle:
        mid = int(box[1] + (box[3] - box[1]) * 0.44)
        draw_text_center(draw, (box[0] + 14, box[1] + 8, box[2] - 14, mid), title, title_font,
                         title_color, max_chars=14)
        draw_text_center(draw, (box[0] + 20, mid - 2, box[2] - 20, box[3] - 8), subtitle,
                         subtitle_font, "#667085", max_chars=21, spacing=5)
    else:
        draw_text_center(draw, box, title, title_font, title_color, max_chars=14)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int],
          color: str = "#2563EB", width: int = 5) -> None:
    draw.line((start, end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    p1 = (end[0] - size * math.cos(angle - 0.55), end[1] - size * math.sin(angle - 0.55))
    p2 = (end[0] - size * math.cos(angle + 0.55), end[1] - size * math.sin(angle + 0.55))
    draw.polygon([end, p1, p2], fill=color)


def canvas(title: str, subtitle: str = "") -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1800, 1050), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), title, font=font(FONT_YAHEI_BOLD, 48), fill="#1E4F9A")
    if subtitle:
        draw.text((72, 112), subtitle, font=font(FONT_YAHEI, 24), fill="#667085")
    draw.line((70, 158, 1730, 158), fill="#2563EB", width=4)
    return image, draw


def generate_diagrams() -> None:
    DIAGRAMS.mkdir(parents=True, exist_ok=True)

    image, draw = canvas("智途 ZhiPath 总体技术架构", "从用户交互到智能体、数据与基础设施的分层实现")
    layers = [
        ("交互体验层", "React 19 / Vite 8 / Three.js / CodeMirror\n学生端、管理端、智能体办公室", "#EAF2FF", "#2563EB"),
        ("业务编排层", "NestJS 11 / JWT / SSE / BullMQ\n岗位、学习、评价、简历、考试、资讯", "#FFF4E8", "#E64A2E"),
        ("智能能力层", "LangGraph + 专业智能体 + DeepSeek\n意图路由、资源生成、审校、规划与匹配", "#ECFDF3", "#2F8F5B"),
        ("数据与基础设施层", "MySQL / Redis / MongoDB / Neo4j / Chroma / MinIO\nSearXNG + Browserless + Docker Compose", "#F4F3FF", "#6941C6"),
    ]
    y = 205
    for i, (name, desc, fill, color) in enumerate(layers):
        box = (165, y, 1635, y + 160)
        rounded_box(draw, box, name, desc, fill, color, color)
        if i < len(layers) - 1:
            arrow(draw, (900, y + 160), (900, y + 205), color)
        y += 210
    image.save(DIAGRAMS / "01-overall-architecture.png")

    image, draw = canvas("岗位驱动的学习成长闭环", "岗位不是终点，而是画像、计划、评价和简历共同迭代的驱动信号")
    nodes = [
        ("真实 JD 获取", "本地岗位 + SearXNG 联网"),
        ("能力匹配", "六因子加权与技能差距"),
        ("主支线计划", "岗位主线 + 自选支线"),
        ("学习与评价", "任务、考试、资源、反馈"),
        ("技能快照", "Commit / Snapshot / Radar"),
        ("简历与投递", "岗位定制与版本管理"),
    ]
    cx, cy, rx, ry = 900, 600, 600, 300
    centers: list[tuple[int, int]] = []
    for i, (title, sub) in enumerate(nodes):
        a = -math.pi / 2 + i * 2 * math.pi / len(nodes)
        centers.append((int(cx + rx * math.cos(a)), int(cy + ry * math.sin(a))))
    for i, center in enumerate(centers):
        nxt = centers[(i + 1) % len(centers)]
        angle = math.atan2(nxt[1] - center[1], nxt[0] - center[0])
        arrow(draw, (int(center[0] + 145 * math.cos(angle)), int(center[1] + 85 * math.sin(angle))),
              (int(nxt[0] - 145 * math.cos(angle)), int(nxt[1] - 85 * math.sin(angle))), "#2563EB", 5)
    for i, ((title, sub), (x, y)) in enumerate(zip(nodes, centers)):
        color = ["#2563EB", "#E64A2E", "#2F8F5B", "#6941C6", "#0E7490", "#B54708"][i]
        rounded_box(draw, (x - 155, y - 90, x + 155, y + 90), title, sub, "#FFFFFF", color, color)
    rounded_box(draw, (710, 500, 1090, 700), "个人成长目标", "职业方向持续校准\n学习证据可追踪", "#FFF9E8", "#E6A817", "#B54708")
    image.save(DIAGRAMS / "02-career-loop.png")

    image, draw = canvas("Git 思想驱动的主支线学习模型", "把学习过程映射为 Branch、Commit、Snapshot、Merge 与 Rollback")
    draw.line((190, 380, 1600, 380), fill="#2563EB", width=8)
    main_points = [(260, "基线画像"), (600, "完成 React"), (990, "通过阶段考试"), (1430, "岗位匹配提升")]
    for x, label in main_points:
        draw.ellipse((x - 22, 358, x + 22, 402), fill="#2563EB", outline="#FFFFFF", width=4)
        draw_text_center(draw, (x - 130, 415, x + 130, 500), label, font(FONT_YAHEI, 22), "#1E4F9A", 10)
    draw.text((155, 295), "main / 岗位主线", font=font(FONT_YAHEI_BOLD, 30), fill="#1E4F9A")
    draw.line((600, 380, 780, 650), fill="#E64A2E", width=7)
    draw.line((780, 650, 1160, 650), fill="#E64A2E", width=7)
    draw.line((1160, 650, 1430, 380), fill="#E64A2E", width=7)
    side_points = [(780, "自选：Three.js"), (1160, "支线成果合并")]
    for x, label in side_points:
        draw.ellipse((x - 22, 628, x + 22, 672), fill="#E64A2E", outline="#FFFFFF", width=4)
        draw_text_center(draw, (x - 150, 690, x + 150, 770), label, font(FONT_YAHEI, 22), "#B42318", 12)
    draw.text((645, 795), "side / 自选支线", font=font(FONT_YAHEI_BOLD, 30), fill="#B42318")
    boxes = [
        ("Commit", "记录一次学习动作与证据"),
        ("Snapshot", "固化技能、雷达与能力指标"),
        ("Compare", "查看前后增量和分支差异"),
        ("Rollback", "非破坏式回到历史快照"),
    ]
    for i, (t, s) in enumerate(boxes):
        x = 170 + i * 405
        rounded_box(draw, (x, 865, x + 330, 1005), t, s, "#FFFFFF", "#94A3B8", "#27303F", 2)
    image.save(DIAGRAMS / "03-git-learning.png")

    image, draw = canvas("多智能体协作与资源生产", "LangGraph 状态机连接意图识别、动作执行、专业智能体与资源台账")
    rounded_box(draw, (90, 420, 340, 610), "用户对话", "自然语言目标\n上下文与页面场景", "#EAF2FF", "#2563EB", "#1E4F9A")
    rounded_box(draw, (440, 260, 760, 450), "意图路由", "IntentRouter\n工具选择与参数抽取", "#FFF4E8", "#E64A2E", "#B42318")
    rounded_box(draw, (440, 585, 760, 775), "LangGraph 引擎", "状态、条件边、重试\n动作与消息聚合", "#ECFDF3", "#2F8F5B", "#166534")
    agents = [("讲义", 880, 220), ("代码", 1180, 220), ("阅读", 1480, 220),
              ("考试", 880, 500), ("评估", 1180, 500), ("岗位", 1480, 500),
              ("路径", 880, 780), ("简历", 1180, 780), ("视频", 1480, 780)]
    for name, x, y in agents:
        rounded_box(draw, (x - 110, y - 70, x + 110, y + 70), f"{name}智能体", "", "#FFFFFF", "#6941C6", "#4A1FB8", 2)
    arrow(draw, (340, 500), (440, 355), "#2563EB")
    arrow(draw, (340, 535), (440, 675), "#2563EB")
    arrow(draw, (760, 355), (880 - 110, 290), "#E64A2E")
    arrow(draw, (760, 675), (880 - 110, 710), "#2F8F5B")
    draw.text((890, 955), "生成结果统一进入 generated_resources_v3，并通过 SSE 回传页面与智能体办公室",
              font=font(FONT_YAHEI, 24), fill="#667085")
    image.save(DIAGRAMS / "04-multi-agent.png")

    image, draw = canvas("联网 JD 获取与透明降级链路", "真实来源优先；联网失败时明确标注 AI 补充，避免来源误导")
    steps = [
        ("输入关键词", "岗位 / 公司 / 技能"),
        ("本地多字段检索", "标题、公司、城市、JD、技能"),
        ("search-stack", "百度 / Bing / 360 / 搜狗"),
        ("LLM 结构化提取", "title / company / skills / URL"),
        ("去重与匹配排序", "来源、技能与匹配度"),
        ("透明展示", "联网岗位 / AI 补充分别标记"),
    ]
    x = 80
    colors = ["#2563EB", "#0E7490", "#6941C6", "#2F8F5B", "#E64A2E", "#B54708"]
    for i, ((title, sub), color) in enumerate(zip(steps, colors)):
        rounded_box(draw, (x, 360, x + 240, 575), title, sub, "#FFFFFF", color, color, 3)
        if i < len(steps) - 1:
            arrow(draw, (x + 240, 468), (x + 290, 468), "#94A3B8", 4)
        x += 290
    rounded_box(draw, (650, 700, 1150, 900), "失败降级", "搜索或提取失败 → DeepSeek 生成推荐\n卡片标记“AI 推荐”，详情提示非真实在招", "#FFF1F0", "#E64A2E", "#B42318", 3)
    arrow(draw, (925, 575), (900, 700), "#E64A2E", 5)
    image.save(DIAGRAMS / "05-jd-search.png")

    image, draw = canvas("岗位匹配六因子模型", "按校招/社招场景动态使用权重，输出总分、差距和投递门槛")
    labels = ["必须技能", "加分技能", "项目经历", "考试成绩", "学习进度", "学习速度"]
    values = [0.30, 0.15, 0.15, 0.20, 0.10, 0.10]
    center = (610, 595)
    radius = 320
    for ring in range(1, 6):
        points = []
        for i in range(6):
            a = -math.pi / 2 + i * math.pi / 3
            points.append((center[0] + radius * ring / 5 * math.cos(a), center[1] + radius * ring / 5 * math.sin(a)))
        draw.polygon(points, outline="#CBD5E1")
    for i, label in enumerate(labels):
        a = -math.pi / 2 + i * math.pi / 3
        end = (center[0] + radius * math.cos(a), center[1] + radius * math.sin(a))
        draw.line((center, end), fill="#CBD5E1", width=2)
        lx = center[0] + (radius + 85) * math.cos(a)
        ly = center[1] + (radius + 55) * math.sin(a)
        draw_text_center(draw, (int(lx - 90), int(ly - 35), int(lx + 90), int(ly + 35)), label,
                         font(FONT_YAHEI, 22), "#27303F", 8)
    sample = [0.78, 0.60, 0.72, 0.80, 0.66, 0.55]
    poly = []
    for i, val in enumerate(sample):
        a = -math.pi / 2 + i * math.pi / 3
        poly.append((center[0] + radius * val * math.cos(a), center[1] + radius * val * math.sin(a)))
    draw.polygon(poly, fill="#93C5FD", outline="#2563EB")
    draw.text((1100, 300), "校招默认权重", font=font(FONT_YAHEI_BOLD, 32), fill="#1E4F9A")
    for i, (label, val) in enumerate(zip(labels, values)):
        y = 375 + i * 85
        draw.text((1100, y), label, font=font(FONT_YAHEI, 24), fill="#27303F")
        draw.rounded_rectangle((1290, y + 2, 1650, y + 35), radius=14, fill="#E2E8F0")
        draw.rounded_rectangle((1290, y + 2, int(1290 + 360 * val / 0.30), y + 35), radius=14, fill="#2563EB")
        draw.text((1665, y), f"{int(val * 100)}%", font=font(FONT_YAHEI_BOLD, 23), fill="#1E4F9A")
    draw.text((1100, 900), "投递判定 = 总分 + 必须技能覆盖率 + 岗位级别附加条件",
              font=font(FONT_YAHEI, 23), fill="#667085")
    image.save(DIAGRAMS / "06-match-model.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color: str, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.font.bold = bold


def set_run_eastasia(run, name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)


def add_field(run, instruction: str) -> None:
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2, fld_char3])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("— ")
    add_field(run, "PAGE")
    paragraph.add_run(" —")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


class DocBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.25)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.45)
        section.right_margin = Cm(2.35)
        section.footer_distance = Cm(0.9)
        add_page_number(section.footer.paragraphs[0])
        self._setup_styles()
        self.figure_no = 0
        self.table_no = 0

    def _setup_styles(self) -> None:
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.first_line_indent = Pt(21)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.space_before = Pt(0)

        specs = {
            "Title": (22, "微软雅黑", BLUE_DARK, True, WD_ALIGN_PARAGRAPH.LEFT),
            "Heading 1": (16, "微软雅黑", BLUE_DARK, True, WD_ALIGN_PARAGRAPH.LEFT),
            "Heading 2": (13, "微软雅黑", BLUE, True, WD_ALIGN_PARAGRAPH.LEFT),
            "Heading 3": (11, "微软雅黑", BLUE, True, WD_ALIGN_PARAGRAPH.LEFT),
        }
        for name, (size, east, color, bold, align) in specs.items():
            style = styles[name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.color.rgb = RGBColor.from_string(color)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
            style.paragraph_format.alignment = align
            style.paragraph_format.first_line_indent = Pt(0)
            style.paragraph_format.space_before = Pt(8 if name != "Title" else 0)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.keep_with_next = True

        if "Caption ZhiPath" not in styles:
            cap = styles.add_style("Caption ZhiPath", WD_STYLE_TYPE.PARAGRAPH)
        else:
            cap = styles["Caption ZhiPath"]
        cap.font.name = "Times New Roman"
        cap.font.size = Pt(9)
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(6)
        cap.paragraph_format.first_line_indent = Pt(0)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def body(self, text: str, bold_prefix: str | None = None) -> None:
        p = self.doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            p.add_run(bold_prefix).bold = True
            p.add_run(text[len(bold_prefix):])
        else:
            p.add_run(text)

    def bullets(self, items: Iterable[str], ordered: bool = False) -> None:
        for i, item in enumerate(items, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.left_indent = Pt(21)
            p.paragraph_format.hanging_indent = Pt(14)
            p.add_run(f"{i}. " if ordered else "• ").bold = ordered
            p.add_run(item)

    def table(self, headers: Sequence[str], rows: Sequence[Sequence[str]], caption: str | None = None,
              widths: Sequence[float] | None = None) -> None:
        self.table_no += 1
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr = table.rows[0]
        set_repeat_table_header(hdr)
        for i, text in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = str(text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, BLUE)
            set_cell_text_color(cell, "FFFFFF", True)
            if widths:
                cell.width = Cm(widths[i])
        for r_index, row in enumerate(rows):
            cells = table.add_row().cells
            for i, text in enumerate(row):
                cells[i].text = str(text)
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if widths:
                    cells[i].width = Cm(widths[i])
                if r_index % 2:
                    set_cell_shading(cells[i], "F8FAFC")
                for paragraph in cells[i].paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)
                        set_run_eastasia(run, "宋体")
        if caption:
            p = self.doc.add_paragraph(style="Caption ZhiPath")
            p.add_run(f"表 {self.table_no}  {caption}")

    def figure(self, path: Path, caption: str, width_cm: float = 15.8) -> None:
        self.figure_no += 1
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        cp = self.doc.add_paragraph(style="Caption ZhiPath")
        cp.add_run(f"图 {self.figure_no}  {caption}")

    def chapter(self, title: str) -> None:
        self.page_break()
        self.heading(title, 1)

    def section_page(self, title: str) -> None:
        self.heading(title, 2)


def add_cover_abstract(b: DocBuilder) -> None:
    p = b.doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.add_run("智途 ZhiPath——基于多智能体与 Git 分支学习模型的\n岗位驱动智能成长平台").bold = True
    p.style = b.doc.styles["Title"]
    p.paragraph_format.line_spacing = 1.15
    sub = b.doc.add_paragraph()
    sub.paragraph_format.first_line_indent = Pt(0)
    r = sub.add_run("作品设计实现方案")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    set_run_eastasia(r, "微软雅黑")
    line = b.doc.add_paragraph()
    line.paragraph_format.first_line_indent = Pt(0)
    line.paragraph_format.space_after = Pt(8)
    run = line.add_run("━" * 48)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    b.heading("摘要", 2)
    abstract = [
        "面向高校学生的软件工程学习与就业准备长期存在三类断裂：岗位要求与学习内容分离，学习过程与能力证据分离，能力画像与简历投递分离。传统平台通常以课程或题库为中心，学习者难以回答“目标岗位缺什么、今天应学什么、学习后能力提升了多少、这些提升如何进入简历”四个连续问题。",
        "智途 ZhiPath 构建了以真实岗位为驱动、以 Git 学习轨迹为证据、以多智能体协作为生产力的智能成长平台。系统将本地岗位库与 SearXNG 联网 JD 获取结合，通过六因子匹配模型识别技能差距；把岗位主线计划与自选支线计划映射为 main/side 分支，每次学习、考试、资源完成和 AI 评价均形成 commit，并生成技能 snapshot、delta 与固定维度能力雷达；讲义、代码、阅读、考试、评估、岗位、路径、简历和视频等专业智能体由 LangGraph 状态机统一编排，生成结果进入可恢复的资源台账。",
        "在交互层，平台提供岗位匹配、学习组合、Git 进度、能力三维可视化、智能体办公室、错题本、考试中心与简历版本管理。联网岗位严格区分真实来源和 AI 补充：真实岗位保留来源链接，联网未命中时的模型推荐以“AI 推荐”标识并说明不代表企业当前招聘，从业务语义上避免来源误导。简历智能体基于个人画像、项目经历、技能证据和目标岗位生成结构化内容，支持预览、分支版本和 PDF 导出。",
        "系统采用 React 19、Vite 8、TypeScript、Three.js、NestJS 11、TypeORM、MySQL、Redis、BullMQ、LangGraph 与 DeepSeek/OpenAI 兼容接口构建，并通过 Docker Compose 集成 MongoDB、Neo4j、Chroma、MinIO、SearXNG 与 Browserless。当前代码验证结果为后端 12 个 Jest 测试套件、37 个测试全部通过，NestJS 与 React/Vite 构建通过，14 个核心界面在 1440×900 浏览器回归中均无控制台错误。",
        "本作品的主要创新在于：提出岗位目标与学习证据联动的职业成长闭环；以 Git 分支思想统一主线、支线、提交、快照、合并和回滚；建立评价记录与技能快照的证据链；以透明降级机制处理联网 JD 不确定性；将智能体运行状态、资源生产和用户学习进度在同一工作空间中可视化。",
    ]
    for text in abstract:
        b.body(text)
    kw = b.doc.add_paragraph()
    kw.paragraph_format.first_line_indent = Pt(0)
    kw.add_run("关键词：").bold = True
    kw.add_run("岗位驱动学习；多智能体协作；Git 学习模型；技能快照；联网 JD；能力画像；智能简历")
    info = b.doc.add_paragraph()
    info.paragraph_format.first_line_indent = Pt(0)
    info.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info.paragraph_format.space_before = Pt(12)
    info.add_run("智途 ZhiPath 项目组  |  2026 年 7 月").font.color.rgb = RGBColor.from_string(GRAY)


def add_toc(b: DocBuilder) -> None:
    b.page_break()
    b.heading("目录", 1)
    p = b.doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    add_field(run, 'TOC \\o "1-3" \\h \\z \\u')
    note = b.doc.add_paragraph()
    note.paragraph_format.first_line_indent = Pt(0)
    note.add_run("注：打开 Word 后目录可通过“更新域”重新计算；交付 PDF 已完成页码更新。").italic = True


def add_chapter_one(b: DocBuilder) -> None:
    b.chapter("第一章 项目概述")
    b.heading("1.1 背景与目标", 2)
    b.heading("1.1.1 项目背景", 3)
    for text in [
        "生成式人工智能正在改变软件工程学习方式，但“能回答问题”不等于“能形成可验证的职业能力”。高校学生往往同时面对课程学习、项目实践、技术选型、求职准备和简历表达，多条任务线相互竞争，传统学习平台很难根据具体岗位目标持续调度。",
        "企业招聘要求以 JD 形式快速变化，同一“前端工程师”岗位可能同时包含 React、TypeScript、工程化、性能优化、跨端开发和业务理解。若学习路径只按教材章节排列，学习者完成课程后仍无法判断与岗位之间的差距；若直接由大模型生成学习建议，又容易出现一次性、不可追踪和缺少证据的问题。",
        "ZhiPath 将岗位数据视为动态目标，将用户技能、项目、考试和学习提交视为能力证据。系统不是简单叠加聊天机器人，而是把岗位检索、匹配分析、计划调度、资源生产、评价沉淀、技能快照和简历生成连接为可回溯的工程系统。",
    ]:
        b.body(text)
    b.heading("1.1.2 国内外产品与研究现状", 3)
    b.body("现有职业教育平台大致分为课程平台、题库测评、求职招聘和通用 AI 助手四类。课程平台内容体系完整，但难以实时响应岗位变化；题库可以验证局部知识，却无法形成持续的项目与技能证据；招聘平台拥有真实岗位，却通常只提供推荐或投递；通用 AI 助手交互灵活，但缺少长期画像、业务状态和可审计动作。")
    b.table(
        ["类别", "代表能力", "主要优势", "对本项目暴露的不足"],
        [
            ["MOOC/课程平台", "课程、章节、作业", "内容体系稳定", "岗位目标与课程进度弱关联"],
            ["题库/测评平台", "刷题、评分、错题", "反馈速度快", "能力证据局限于答题结果"],
            ["招聘平台", "JD、推荐、投递", "岗位数据真实", "缺少学习闭环与技能演进"],
            ["通用 AI 助手", "问答、生成、解释", "自然语言交互灵活", "状态短期、结果难追踪、易出现来源误导"],
            ["ZhiPath", "岗位—学习—评价—简历", "目标、过程与证据统一", "需要持续扩展真实 JD 与评价样本"],
        ],
        "相关产品能力对比",
        [2.8, 3.2, 3.1, 6.2],
    )
    b.heading("1.1.3 痛点分析", 3)
    b.bullets([
        "目标模糊：学习者知道要“学前端/后端”，但缺少可计算的岗位要求和达成门槛。",
        "计划混乱：岗位必修内容与个人兴趣混在同一清单，优先级和资源投入无法解释。",
        "过程不可追踪：学习动作完成后只剩进度百分比，无法回答何时、因何、提升了什么。",
        "评价碎片化：考试、速测、AI 评估和项目成果分散，无法统一作用于技能画像。",
        "信息来源不透明：联网搜索失败后若直接使用模型生成结果，用户容易误认为是真实招聘。",
        "表达断裂：画像和项目经历无法自动转化为针对岗位的简历版本。",
    ], ordered=True)
    b.heading("1.1.4 项目目标", 3)
    b.body("技术目标是建立可扩展的多智能体业务编排、统一评价主线和 Git 学习数据模型；业务目标是让学生从一个目标岗位出发，获得明确差距、可执行计划、连续能力证据和可投递成果；体验目标是让复杂能力以安静、可扫描的工作台呈现，避免把 AI 能力包装成不可解释的黑箱。")

    b.section_page("1.2 功能概述")
    b.table(
        ["功能域", "核心能力", "关键产物"],
        [
            ["AI 助教", "意图识别、连续对话、智能体动作", "消息、工具调用、生成资源"],
            ["学习组合", "岗位主线、自选支线、阶段与每日任务", "计划、阶段、任务、工时预算"],
            ["Git 学习", "分支、提交、快照、比较、合并、回滚", "Commit、Snapshot、Delta、Radar"],
            ["岗位匹配", "本地搜索、联网 JD、六因子匹配", "岗位卡片、差距、投递门槛"],
            ["考试评价", "考试、速测、错题、AI 评价", "Attempt、Result、Evidence、Impact"],
            ["能力画像", "技能清单、标签、固定维度雷达、3D 图谱", "能力指标、技能节点、关系图"],
            ["简历管理", "岗位定制、分支版本、预览与导出", "结构化简历、版本、PDF"],
            ["智能体办公室", "任务状态、工位、资源与运行事件", "任务台账、资源台账、SSE 状态"],
            ["管理后台", "用户、企业、岗位、考试、资讯管理", "运营数据与配置"],
        ],
        "系统功能模块",
        [2.4, 6.5, 6.0],
    )
    b.figure(DIAGRAMS / "02-career-loop.png", "岗位驱动的学习成长闭环")

    b.heading("1.3 典型用户流程", 2)
    b.heading("1.3.1 首次使用流程", 3)
    b.body("新用户完成注册后进入画像引导，填写学校、专业、年级、职业方向、每日可投入时间、技能和项目。系统以这些信息建立初始技能基线，并引导用户选择本地岗位或输入岗位关键词。确定目标岗位后，匹配服务计算技能差距，规划服务据此创建 main 主线计划，任务调度器再按照每日工时生成可执行任务。")
    b.body("首次使用流程强调“先建立可计算目标，再生成学习内容”。如果用户暂时没有明确岗位，可以使用岗位浏览和 AI 助教澄清方向；如果已经拥有 JD，可以通过岗位搜索或管理员岗位库直接选择。系统不会仅凭专业名称生成完整规划，而是将画像、目标岗位、技能差距和时间约束共同作为输入。")
    b.heading("1.3.2 日常学习流程", 3)
    b.body("用户每天从首页进入今日任务，可阅读讲义、完成代码练习、参加速测或正式考试。每个动作完成后由 LearningCommitService 形成 commit，SkillSnapshotService 生成新快照和能力差值，评价系统记录得分、证据和影响。页面通过 SSE 接收资源和智能体状态，无需停留在聊天页面等待长任务。")
    b.heading("1.3.3 求职准备流程", 3)
    b.body("当目标岗位或个人能力发生变化时，系统重新计算匹配结果，给出已匹配技能、缺失技能和投递门槛。用户可继续学习，也可基于当前证据生成岗位定制简历。简历版本与岗位关联，保留 AI 生成、人工修改和分支历史；真实联网岗位可跳转来源页面，AI 补充岗位仅作为方向参考。")

    b.heading("1.4 系统边界与角色", 2)
    b.table(
        ["角色", "核心权限", "明确边界"],
        [
            ["学生用户", "画像、学习、考试、岗位、简历、智能体", "不能修改企业和公共岗位基础数据"],
            ["管理员", "用户、企业、岗位、考试、资讯与配置管理", "不能替代学生完成学习评价"],
            ["专业智能体", "生成资源、规划、评估、提取和整理", "不得伪造用户经历或真实招聘状态"],
            ["外部搜索服务", "提供公开网页候选结果", "结果时效与真实性以来源页为准"],
            ["外部模型服务", "语言理解与内容生成", "不直接持有业务事务和最终事实"],
        ],
        "系统角色与权限边界",
    )
    b.body("系统将 AI 定位为“受业务规则约束的协作者”，而不是拥有无限权限的代理。模型不能直接修改分支 head、技能掌握度、投递记录或管理员数据；所有具有业务后果的操作都由后端服务校验，并留下用户、时间、动作和证据。")

    b.heading("1.5 可行性分析", 2)
    b.heading("1.5.1 技术可行性", 3)
    b.body("React、NestJS、MySQL、Redis 和 Docker 均为成熟技术，LangGraph 为多步骤智能体提供状态机抽象，SearXNG 与 Browserless 支持可替换的联网搜索和网页抓取。项目已通过完整构建与单元测试，并在本地部署环境完成学生端、管理端和搜索栈联调。")
    b.heading("1.5.2 业务可行性", 3)
    b.body("学生的岗位选择、技能差距、每日任务、考试评价和简历准备具有明确的先后关系，适合以统一数据链路管理。系统允许从一个功能切入逐步使用：用户可只使用岗位匹配，也可在此基础上启用学习计划、评价和简历，不要求一次性完成全部数据。")
    b.heading("1.5.3 运维可行性", 3)
    b.body("核心事务依赖 MySQL，Redis、MongoDB、Neo4j、Chroma 和 MinIO 可按 profile 分阶段启用。外部 AI、搜索、地图和 TTS 均通过环境变量和服务适配器配置；当可选中间件不可用时，核心登录、本地岗位、计划和基础数据仍可运行。")


def add_chapter_two(b: DocBuilder) -> None:
    b.chapter("第二章 核心技术")
    b.heading("2.1 智能学习者画像与能力可视化", 2)
    b.heading("2.1.1 内容概述", 3)
    b.body("画像模块以用户基本信息、技能条目、项目经历、学习快照和评价结果为输入。与仅展示技能标签的画像不同，ZhiPath 同时维护“技能事实”和“能力维度”两类视图：前者回答会什么、掌握到何种程度，后者回答综合能力结构是否均衡。")
    b.heading("2.1.2 核心技术", 3)
    b.bullets([
        "技能事实层：user_skills_v3 保存技能名称、掌握度、证据来源与更新时间。",
        "快照层：skill_snapshots_v3 固化某次 commit 后的完整技能集、固定维度雷达和能力指标。",
        "可视化层：React Three Fiber、Three.js 与 drei 构建可旋转缩放的能力关系图；中心球表示个人能力，技能节点按类别和掌握度布局。",
        "计算层：SkillSnapshotService 计算 snapshot、delta、radar 和 ability metrics，避免前端自行拼接历史数据。",
        "版本层：画像变化绑定 commitId，能力提升可以回溯到考试、资源、任务或人工操作。",
    ])
    b.table(
        ["能力维度", "主要证据", "更新时机", "展示方式"],
        [
            ["技术基础", "语言、框架、课程能力", "技能提交/考试", "雷达轴 + 技能节点"],
            ["工程实践", "项目、代码任务、工程化技能", "项目/代码资源完成", "雷达轴 + 项目关联"],
            ["学习执行", "任务完成、学习时长、连续性", "每日任务与阶段完成", "指标卡 + 趋势"],
            ["问题解决", "考试、速测、错题修复", "评价结果落库", "雷达轴 + 差距"],
            ["职业匹配", "岗位六因子匹配", "目标岗位或画像变化", "匹配分 + 岗位节点"],
            ["表达协作", "项目描述、简历证据、对话", "简历/评价更新", "能力节点 + 简历版本"],
        ],
        "能力维度与证据映射",
    )
    b.heading("2.1.3 创新点和优势", 3)
    b.body("画像不是独立问卷结果，而是 Git 学习链上的派生状态。每个指标都能找到提交、评价或岗位匹配来源；三维图谱承担探索和关系理解，二维雷达承担稳定比较，两种视图各司其职。")

    b.section_page("2.2 多智能体协作引擎")
    b.body("多智能体模块采用“统一编排、专业执行、资源落账”的结构。聊天入口首先由 IntentRouter 判断用户意图，LangGraphEngine 维护消息、动作、资源和状态，ActionExecutor 调用专业服务。智能体结果不是只写入当前对话，而是统一保存到 generated_resources_v3，页面切换或会话恢复后仍可访问。")
    b.figure(DIAGRAMS / "04-multi-agent.png", "多智能体协作与资源生产架构")
    b.table(
        ["智能体", "职责", "典型输入", "输出"],
        [
            ["Orchestrator", "拆解目标与调度动作", "对话上下文、页面场景", "动作序列"],
            ["ProfileAgent", "整理画像与技能", "用户资料、历史证据", "画像建议"],
            ["PathAgent", "生成或调整学习路径", "目标岗位、差距、工时", "阶段与技能计划"],
            ["LectureAgent", "生成结构化讲义", "技能、难度、薄弱点", "Markdown 讲义"],
            ["CodeAgent", "生成代码练习与解释", "技能、语言、场景", "代码资源"],
            ["ReadingAgent", "生成延伸阅读", "技能与目标", "阅读材料"],
            ["ExamAgent", "生成考试内容", "技能、难度、题型", "题目与答案"],
            ["ReviewerAgent", "审校答案与内容", "生成资源、评分规则", "审校与修正"],
            ["JdParserAgent", "结构化岗位描述", "JD 文本或搜索摘要", "岗位字段"],
            ["ResumeAgent", "生成岗位定制简历", "画像、项目、岗位", "简历 JSON"],
            ["VideoAgent", "组织脚本、语音与视频", "知识主题、讲解目标", "教学视频"],
        ],
        "专业智能体职责",
    )
    b.heading("2.2.1 状态机与容错", 3)
    b.body("LangGraph 以显式 State 保存消息、意图、动作、资源和错误，条件边决定继续执行、请求澄清或结束。外部服务失败时，动作执行器保留可用的局部结果；长任务通过 BullMQ 排队，并使用 SSE 广播状态。生成 JSON 统一经过 repair/extract 流程，兼容代码块、包装对象和截断输出。")

    b.section_page("2.3 岗位检索、JD 获取与匹配")
    b.body("岗位模块同时支持本地结构化检索与联网检索。本地岗位对标题、公司、城市、薪资、JD、必备技能和加分技能进行多字段匹配；联网模式调用独立 search-stack，由 SearXNG 聚合国内可访问引擎，再由 LLM 将搜索结果提取为岗位卡片。")
    b.figure(DIAGRAMS / "05-jd-search.png", "联网 JD 获取与透明降级链路")
    b.heading("2.3.1 透明来源模型", 3)
    b.body("系统为在线卡片增加 origin 字段，区分 web 与 ai_generated。真实来源保留 URL 和 host，可在内部详情页查看摘要并跳转原岗位；模型补充卡片显示“AI 推荐”，详情页明确提示“并不代表企业当前正在招聘”，并隐藏无效投递操作。该设计将技术降级显式转换为业务状态，而不是用成功外观掩盖失败。")
    b.figure(DIAGRAMS / "06-match-model.png", "岗位匹配六因子模型")
    b.heading("2.3.2 六因子匹配", 3)
    b.body("匹配服务综合必须技能、加分技能、项目经历、考试成绩、学习进度和学习速度。校招场景关注考试与学习证据，社招场景提升项目和必须技能权重。除总分外，服务返回 matched、missing、gapAnalysis、factor breakdown 和 canApply，避免把单一分数当作完整判断。")

    b.section_page("2.4 Git 思想驱动的主支线学习")
    b.body("ZhiPath 将岗位目标对应的学习计划定义为 main 主线，将兴趣、补充技能和临时探索定义为 side 支线。主线强调目标达成和稳定排期，支线允许独立推进；支线成果经过比较后可以 merge 到主线。每次学习动作形成 commit，commit 关联 snapshot，从而获得类似软件版本控制的可解释历史。")
    b.figure(DIAGRAMS / "03-git-learning.png", "Git 思想驱动的主支线学习模型")
    b.table(
        ["Git 概念", "学习语义", "系统行为"],
        [
            ["Branch", "一条独立学习目标线", "main 岗位主线、side 自选支线"],
            ["Commit", "一次可解释学习动作", "记录技能、证据、消息、父提交"],
            ["Snapshot", "提交后的完整能力状态", "保存技能、雷达、指标与 head"],
            ["Delta", "前后能力变化", "新增/提升/下降技能与雷达差值"],
            ["Compare", "两条分支或快照差异", "比较共同技能、独有技能和能力差"],
            ["Merge", "把支线成果并入主线", "合并技能与证据，形成新提交"],
            ["Rollback", "回到历史能力状态", "非破坏式更新分支 head 并保留历史"],
        ],
        "Git 概念与学习语义映射",
    )
    b.heading("2.4.1 排期规则", 3)
    b.body("主线计划可以处于 active、paused、archived 状态，同一用户只允许一个 active 主线；支线可并存。每日工时由 dailyHours 与 mainRatio 分配，主线优先获得稳定预算，自选任务只进入 side 计划，避免额外学习内容污染岗位达成进度。")

    b.section_page("2.5 统一评价主线与自适应测评")
    b.body("评价系统通过 EvaluationAttempt、EvaluationResult、EvaluationDimensionScore、EvaluationEvidence 与 EvaluationImpact 五类实体统一考试、速测、AI 评估和过程评价。Attempt 描述一次评价会话，Result 给出结论，DimensionScore 保存分维度得分，Evidence 记录答案、代码或任务证据，Impact 说明本次评价如何影响技能、计划和匹配。")
    b.table(
        ["评价来源", "输入", "输出", "对学习 Git 的影响"],
        [
            ["正式考试", "题目与答案", "总分、错题、知识点", "生成考试 commit 与 snapshot"],
            ["快速测验", "目标技能、少量题", "即时掌握度", "生成轻量 commit"],
            ["AI 评估", "项目/代码/回答", "Rubric 分维度结果", "评价证据关联提交"],
            ["任务完成", "资源或阶段状态", "过程性完成证据", "推进分支 head"],
            ["错题复习", "错误记录与再答", "修复状态", "更新薄弱点与能力差"],
        ],
        "统一评价入口",
    )
    b.body("考试模块支持题目生成、作答、评分、错题分析和评价落库。自适应策略可根据连续正确、连续错误、知识点薄弱和当前掌握度调节难度。用户不只看到一次分数，还能在错题本、画像和 Git 进度页观察评价的后续影响。")

    b.section_page("2.6 智能简历生成与版本管理")
    b.body("简历智能体以结构化画像、技能证据、项目经历、教育信息和目标岗位为上下文，生成 summary、skills、projects、education、experience 等字段。生成过程禁止虚构缺失经历；证据不足时使用克制表述，并将岗位关键词自然映射到已有项目和技能。")
    b.bullets([
        "岗位定制：同一用户可针对不同岗位生成不同简历版本。",
        "分支版本：保留原简历并创建分支，便于比较和回退。",
        "可编辑预览：生成结果进入编辑器而不是直接定稿。",
        "稳定导出：后端 PDF 服务统一字体、分页和中文渲染。",
        "能力证据引用：技能来源于画像、项目和评价，不以模型猜测补齐。",
    ])
    b.table(
        ["阶段", "输入", "处理", "输出"],
        [
            ["数据准备", "画像、技能、项目、岗位", "规范化与缺失检查", "ResumeContext"],
            ["内容生成", "ResumeContext", "DeepSeek 分段生成", "结构化简历 JSON"],
            ["质量校验", "简历 JSON", "字段完整性、事实与长度", "可编辑版本"],
            ["版本管理", "用户修改/目标变化", "更新或分支", "ResumeVersion"],
            ["导出", "最终版本", "模板排版与 PDF 生成", "可投递文件"],
        ],
        "智能简历生成流程",
    )

    b.section_page("2.7 智能体办公室与多模态资源")
    b.body("智能体办公室把后台任务可视化为工位、状态和资源流。讲义、代码、阅读、考试、视频等智能体拥有独立形象和工位，任务开始、进度变化、完成或失败通过 SSE 实时同步。右侧台账展示正在执行的任务和已经生成的资源，使“AI 正在做什么”成为可观察状态。")
    b.body("视频链路组合脚本规划、TTS、可视化渲染和文件输出；知识资源支持 Markdown、代码高亮、Mermaid 与图谱；三维能力图使用 Three.js。不同模态最终统一关联用户、技能、计划、分支和资源类型，避免生成文件成为孤立资产。")
    b.heading("2.7.1 图解资源生成", 3)
    b.body("当用户提出“生成某技术的流程图、架构图或记忆图”时，意图路由把请求识别为图解资源任务。智能体先抽取主题、关键节点、关系方向和解释粒度，再生成 Mermaid 或结构化图数据；前端在聊天消息中渲染为图解卡片，并提供源码查看入口。生成记录进入资源台账，关联原始会话、用户、技能与时间。")
    b.body("以智能体记忆工程流程图为例，资源不仅包含节点名称，还需要表达感知输入、记忆重要性评估、长期/短期存储、检索、衰减、上下文整合、行为决策和反馈回路。系统把这种关系型内容交给图解模态，而不是强行输出长段文字，从而提升结构理解效率。")
    b.heading("2.7.2 教学视频生成", 3)
    b.body("视频资源采用“主题分析—脚本分镜—语音合成—画面渲染—文件入库”流水线。VideoAgent 生成讲解目标和分段脚本，TTS 服务将每段文字转换为音频，视频渲染器根据代码、标题和动画模板合成画面，最终输出 MP4 并记录资源状态。任何阶段失败都会保留任务日志和已完成中间产物，便于重试。")
    b.body("Docker 教学视频示例展示了代码片段、命令与字幕同步播放。视频生成完成后，用户可继续要求路径规划师围绕同一主题生成专项学习计划。由于会话状态保留了 Docker 主题、目标岗位和用户技能，后续计划无需重复输入上下文，体现多模态资源与规划智能体之间的协作。")
    b.heading("2.7.3 多模态资源一致性", 3)
    b.table(
        ["资源类型", "主要生成器", "核心结构", "前端呈现", "持久化信息"],
        [
            ["知识讲义", "LectureAgent", "Markdown 章节", "阅读卡片/详情页", "标题、正文、技能、来源"],
            ["代码示例", "CodeAgent", "语言、代码、解释", "代码编辑器/高亮", "代码、语言、测试说明"],
            ["流程图", "ResourceAgent", "节点、边、Mermaid", "可缩放图解卡片", "源码、预览、主题"],
            ["教学视频", "VideoAgent", "脚本、音频、画面", "内嵌播放器", "文件 URL、时长、状态"],
            ["学习计划", "PathAgent", "阶段、技能、工时", "主支线计划详情", "计划、分支、任务"],
            ["考试资源", "ExamAgent", "题目、答案、知识点", "考试中心", "题库、记录、评价"],
        ],
        "多模态资源的统一结构",
    )

    b.heading("2.8 可靠性、安全与可观测性", 2)
    b.heading("2.8.1 长任务可靠性", 3)
    b.body("讲义、视频、简历和批量资源生成可能超过普通 HTTP 请求时限。系统使用 BullMQ 将任务写入 Redis 队列，保存 waiting、active、completed 和 failed 状态；任务包含 userId、agentType、params、priority 和 groupId，支持按任务或任务组取消。队列超时根据智能体类型和输入复杂度计算，避免用同一阈值处理不同任务。")
    b.heading("2.8.2 缓存与幂等", 3)
    b.body("搜索结果、公司简介和抓取正文设置有限 TTL，减少重复外部调用；学习 commit、评价和资源生成使用业务标识避免重复落库。页面请求通过序列号丢弃过时响应，防止用户快速切换筛选时旧请求覆盖新状态。对于联网岗位，客户端会话缓存只用于临时详情恢复，不替代服务器事实数据。")
    b.heading("2.8.3 身份与密钥", 3)
    b.body("身份认证采用 JWT，后端 Guard 从 Authorization Bearer 或 SSE query token 中解析身份，RolesGuard 限制管理接口。模型、搜索、地图和 TTS 密钥保存在后端环境变量，浏览器只使用地图 Web 端公开配置。生产部署应更换默认 JWT Secret、轮换已暴露密钥，并限制数据库和中间件端口只对本机或内网开放。")
    b.heading("2.8.4 可观测性", 3)
    b.body("服务使用 NestJS Logger、队列统计、SSE 事件和资源台账形成四层观测：接口层记录请求失败，业务层记录降级原因，任务层记录执行状态与耗时，用户层展示可理解的加载、成功、失败和来源状态。联网检索不把“返回数组”视为唯一成功标准，还检查 provider、来源 URL、真实岗位数量和 AI 补充数量。")

    b.heading("2.9 业务约束与异常处理", 2)
    b.table(
        ["异常场景", "系统处理", "用户可见状态", "数据保护"],
        [
            ["模型超时", "终止请求并使用模板/局部结果", "生成失败或降级提示", "不写入空资源"],
            ["搜索引擎超时", "切换其他国内引擎", "继续返回可用来源", "缓存键包含查询和模式"],
            ["联网无真实岗位", "使用 AI 补充推荐", "明确显示 AI 推荐", "不允许虚假投递"],
            ["重复点击搜索", "请求序列号只接收最新响应", "保持当前加载状态", "旧响应不覆盖新数据"],
            ["支线内容误入主线", "服务层校验 planType", "返回业务错误", "主线排期不被污染"],
            ["回滚历史提交", "非破坏式移动 head", "显示目标快照", "历史 commit 保留"],
            ["资源任务失败", "队列记录 failed 与错误", "办公室显示失败", "已完成资源不丢失"],
        ],
        "主要异常场景与处理策略",
    )


def add_chapter_three(b: DocBuilder) -> None:
    b.chapter("第三章 系统架构与实现")
    b.heading("3.1 项目整体架构", 2)
    b.figure(DIAGRAMS / "01-overall-architecture.png", "智途 ZhiPath 总体技术架构")
    b.body("系统采用前后端分离与中间件服务化结构。React 客户端通过统一 Axios Client 访问 /api，JWT 拦截器自动注入令牌；SSE 用于智能体进度、资源完成和通知事件；NestJS 模块承载业务边界，TypeORM 管理结构化数据，队列与缓存削峰长任务；外部 AI、搜索、地图和 TTS 均通过服务适配器封装。")
    b.table(
        ["层次", "组件", "职责"],
        [
            ["客户端", "React、Router、Zustand", "页面、会话状态、交互与可视化"],
            ["接口层", "NestJS Controller、Guard", "REST/SSE、鉴权、参数与响应"],
            ["业务层", "Jobs/Learning/Evaluation/Resume", "领域规则与事务编排"],
            ["智能体层", "LangGraph、Agent Services", "意图、生成、审校、规划"],
            ["任务层", "BullMQ、Redis", "长任务、重试、取消、状态"],
            ["数据层", "MySQL、MongoDB、Neo4j、Chroma", "事务、文档、图、向量"],
            ["对象层", "MinIO", "PDF、音频、视频与附件"],
            ["外部服务", "DeepSeek、SearXNG、Browserless、高德", "模型、搜索、抓取与地图"],
        ],
        "系统分层与职责",
    )

    b.section_page("3.2 前端架构与设计系统")
    b.body("前端以用户任务流为页面边界：AI 助教负责自然语言入口，主页负责态势总览，学习路径负责计划执行，岗位负责目标与匹配，考试和错题本负责评价，画像负责能力解释，智能体办公室负责运行透明。管理端沿用同一手绘设计语言，但以淡红侧栏区分角色。")
    b.bullets([
        "布局：固定侧边导航，主内容区域独立滚动；桌面端优先保证扫描密度。",
        "组件：按钮、标签、画布、卡片、分段控件、进度条和表格共享 hand-draw.css 令牌。",
        "可视化：Three.js 负责能力图与办公室，普通业务图表保持二维、可打印。",
        "交互状态：加载、成功、失败、空结果和降级来源都使用明确文案，不以空白代替状态。",
        "响应式：导航在窄屏折叠，卡片网格从三列转为两列和单列。",
    ])

    ui_pages = [
        ("3.2.1 落地页界面", "01-landing.png", "落地页：职业方向与智能成长入口", ["首屏直接传达岗位驱动价值", "登录注册与产品叙事在同一页面完成", "暖色纸张与手绘线框形成统一品牌识别"]),
        ("3.2.2 首页仪表盘界面", "02-dashboard.png", "学生首页：学习与求职态势总览", ["聚合技能、学习时长、考试和岗位指标", "主计划阶段与今日任务并排展示", "快捷入口减少高频操作层级"]),
        ("3.2.3 AI 助教界面", "03-chat.png", "AI 助教：对话、智能体与资源工作区", ["会话历史、消息流和智能体面板三栏布局", "工具调用结果以内嵌资源卡呈现", "资源台账避免切页后结果丢失"]),
        ("3.2.4 学习组合界面", "04-learning.png", "学习组合：岗位主线与自选支线", ["计划组合总览与当前计划详情共存", "阶段、技能、资源和工时在同一上下文", "主线/支线语义与 Git 分支一致"]),
        ("3.2.5 岗位匹配界面", "05-jobs.png", "岗位匹配：本地、混合与联网搜索", ["模式切换、关键词搜索和筛选分离", "卡片同时展示来源、薪资、技能与匹配度", "联网未命中与 AI 补充数量明确区分"]),
        ("3.2.6 岗位详情界面", "06-job-detail.png", "岗位详情：要求、公司、位置与六因子匹配", ["岗位事实与匹配分析左右分栏", "公司简介由 flash 模型克制整理", "高德地图支持点击激活和滚轮缩放"]),
        ("3.2.7 用户画像界面", "07-profile.png", "用户画像：技能清单与能力可视化", ["基本资料、技能标签和能力图分区组织", "能力节点来自真实技能清单", "三维节点支持旋转、缩放与聚焦"]),
        ("3.2.8 Git 学习进度界面", "08-progress.png", "Git 学习系统：分支、提交与雷达对比", ["分支选择与 commit 时间线可追踪", "雷达图展示当前能力结构", "比较、合并和回滚均以非破坏方式执行"]),
        ("3.2.9 简历管理界面", "09-resume.png", "简历管理：生成、分支、编辑与导出", ["简历版本与目标岗位关联", "支持 AI 生成和手动修订", "预览与 PDF 导出使用同一结构化数据"]),
        ("3.2.10 考试中心界面", "10-exams.png", "考试中心：考试列表、成绩与状态", ["考试统计与列表同屏", "得分、通过状态和操作清晰", "考试结果进入统一评价主线"]),
        ("3.2.11 错题本界面", "11-wrong-answers.png", "错题本：薄弱知识点回收与复习", ["按技能与题目状态组织错误记录", "支持展开分析与再次练习", "修复结果可更新技能快照"]),
        ("3.2.12 智能体办公室界面", "12-agent-office.png", "智能体办公室：任务与资源运行可视化", ["工位展示智能体状态与任务分配", "右侧任务台账和资源台账同步", "SSE 事件驱动实时状态变化"]),
        ("3.2.13 多模态图解资源界面", "15-multimodal-memory-workflow.png", "多模态资源：智能体记忆工程工作流程图", ["用户通过自然语言请求生成技术流程图", "图解资源以内嵌卡片展示并支持查看源码", "右侧智能体办公室同步显示角色和资源数量"]),
        ("3.2.14 视频与计划联动界面", "16-multimodal-video-learning-path.png", "多模态资源：Docker 教学视频与专项学习计划", ["视频智能体生成可直接播放的 Docker 教学视频", "路径规划师根据同一目标生成专项分阶段计划", "视频、计划和对话保留在同一任务上下文中"]),
        ("3.2.15 管理看板界面", "13-admin-dashboard.png", "管理看板：业务规模与运营状态", ["用户、企业、岗位、考试等指标汇总", "待办与最新动态形成运营入口", "淡红侧栏区分管理角色"]),
        ("3.2.16 岗位管理界面", "14-admin-jobs.png", "管理端岗位列表与维护", ["岗位、企业、技能、级别和薪资集中管理", "支持检索、新增、编辑与状态维护", "为本地岗位库和匹配服务提供基础数据"]),
    ]
    for title, filename, caption, notes in ui_pages:
        b.page_break()
        b.heading(title, 2)
        b.figure(SCREENSHOTS / filename, caption, 16.2)
        b.bullets(notes)

    b.section_page("3.3 后端服务集成")
    b.heading("3.3.1 数据库与实体设计", 3)
    b.body("MySQL 保存用户、学生、企业、岗位、计划、任务、技能、考试、简历、分支、提交、快照、评价和资源等核心事务数据。实体采用 v3 表名隔离旧模型，学习 Git 与评价主线通过 userId、branchId、commitId 和 snapshotId 连接。")
    b.table(
        ["领域", "核心实体", "主要关系"],
        [
            ["身份画像", "User、Student、UserSkill", "User 1:1 Student，User 1:N Skill"],
            ["岗位企业", "Enterprise、JobPosition、JobApplication", "Enterprise 1:N Job，User N:M Job"],
            ["学习计划", "LearningPlan、LearningTask、LearningSession", "Plan 1:N Task/Session"],
            ["Git 学习", "LearningBranch、LearningCommit、SkillSnapshotV3", "Branch 1:N Commit，Commit 1:1 Snapshot"],
            ["评价", "EvaluationAttempt/Result/Evidence/Impact", "Attempt 聚合结果、证据与影响"],
            ["资源", "GeneratedResource、AgentTask", "任务生成资源并关联技能/计划"],
            ["考试", "Exam、ExamRecord、WrongAnswer", "考试记录进入评价和 commit"],
            ["简历", "Resume、ResumeVersion", "用户多版本、目标岗位关联"],
        ],
        "主要数据实体",
    )
    b.heading("3.3.2 API 组织", 3)
    b.body("Controller 统一位于 /api 前缀，AuthGuard 校验 Bearer JWT，RolesGuard 区分 student 与 admin。分页接口返回 data、total、page、pageSize，业务接口返回 code、message、data；长任务通过 SSE 和队列状态接口补充。")
    b.table(
        ["业务", "代表 API", "说明"],
        [
            ["岗位", "GET /user/jobs", "本地/混合/联网检索与分页"],
            ["岗位详情", "GET /user/jobs/:id/company-context", "公司简介与地理位置"],
            ["学习计划", "GET /user/learning-paths", "主线、支线与阶段数据"],
            ["Git 分支", "GET /user/git/branches", "分支列表与 head"],
            ["提交", "POST /user/git/branches/:id/commit", "形成学习提交与快照"],
            ["比较", "GET /user/git/snapshots/compare", "能力前后差异"],
            ["简历", "POST /user/resumes/generate", "生成岗位定制简历"],
            ["智能体", "POST /user/agents/run", "提交专业智能体任务"],
            ["事件", "GET /user/events/stream", "SSE 实时事件"],
        ],
        "代表性 API",
    )

    b.section_page("3.4 技术栈详解")
    b.table(
        ["类别", "技术", "采用原因"],
        [
            ["前端框架", "React 19 + TypeScript", "组件化、类型安全与成熟生态"],
            ["构建工具", "Vite 8", "快速开发、现代打包与模块热更新"],
            ["状态管理", "Zustand", "轻量会话状态与持久化"],
            ["三维可视化", "Three.js / R3F / drei", "能力图与办公室交互"],
            ["后端框架", "NestJS 11", "模块化、依赖注入、Guard 与测试能力"],
            ["ORM", "TypeORM", "实体建模、查询构造器与迁移兼容"],
            ["智能编排", "LangGraph", "显式状态、条件边与可恢复工作流"],
            ["任务队列", "BullMQ + Redis", "长任务排队、重试、取消与状态"],
            ["事务数据", "MySQL 8", "核心业务一致性"],
            ["文档与图数据", "MongoDB / Neo4j / Chroma", "非结构化、关系和语义检索"],
            ["对象存储", "MinIO", "音视频、PDF 与附件"],
            ["联网搜索", "SearXNG + Browserless", "多引擎检索与动态页面抓取"],
            ["模型服务", "DeepSeek / OpenAI-compatible / Ollama", "分级模型与本地降级"],
        ],
        "技术栈与选型",
    )
    b.body("模型分级策略将低延迟整理、路由和摘要交给 flash 模型，将复杂推理、出题和长内容交给 pro 模型。所有密钥通过环境变量注入，客户端不保存服务端密钥；地图 Web Key 与安全密钥分别配置。")

    b.heading("3.5 部署架构与网络配置", 2)
    b.body("开发环境由前端 5173 端口、后端 3000 端口、MySQL 3307 映射端口和独立 search-stack 组成。Vite 代理把 /api 转发到 NestJS；后端通过 127.0.0.1 访问数据库和搜索代理；SearXNG、Redis 与 Browserless 位于 Docker 网络内部，只将 search-proxy 的 17080 端口绑定到宿主机。")
    b.table(
        ["服务", "宿主机端口", "网络角色", "健康检查"],
        [
            ["Vite Frontend", "5173", "浏览器入口与开发热更新", "页面与构建"],
            ["NestJS Backend", "3000", "业务 API、SSE 与任务入口", "HTTP /api"],
            ["MySQL", "3307→3306", "事务数据", "连接与查询"],
            ["search-proxy", "17080", "统一搜索/抓取 API", "/health + API Key"],
            ["SearXNG", "容器内 8080", "多引擎聚合", "/healthz"],
            ["search-redis", "容器内 6379", "搜索缓存与限流", "redis ping"],
            ["Browserless", "容器内 3000", "动态网页抓取", "HTTP health"],
            ["MinIO", "9000/9001", "对象与管理控制台", "bucket 检查"],
        ],
        "本地部署端口与职责",
    )
    b.body("中国境内部署的 SearXNG 默认启用百度、Bing、360 搜索和搜狗，并配置阿里云与腾讯公共 DNS。需要访问境外引擎时，容器代理使用 host.docker.internal 指向 Windows 宿主机代理，不能使用 127.0.0.1，因为容器内的回环地址只代表容器自身。")

    b.heading("3.6 代码组织与领域边界", 2)
    b.body("后端 modules 目录按接口业务域组织，services 目录承载跨模块的智能体、匹配、分支、提交、快照、评价、资源和外部服务；entities 目录定义数据库模型。前端 pages/user 与 pages/admin 对应角色页面，components 保存复用组件和三维场景，api 统一封装请求，stores 管理会话与聊天状态。")
    b.table(
        ["代码区域", "代表文件", "职责边界"],
        [
            ["modules/jobs", "jobs.service.ts", "岗位检索、详情、公司上下文与投递"],
            ["modules/learning-paths", "learning-paths.service.ts", "主支线计划、阶段和资源调度"],
            ["modules/git-learning", "git-learning.controller.ts", "分支、提交、快照、比较与合并接口"],
            ["services/branch", "branch.service.ts", "分支业务规则、回滚与合并"],
            ["services/learning-commit", "learning-commit.service.ts", "学习动作提交与证据组织"],
            ["services/skill-snapshot", "skill-snapshot.service.ts", "技能快照、delta、radar、metrics"],
            ["modules/chat", "langgraph-engine.service.ts", "对话状态机和动作编排"],
            ["services/agents", "*-agent.service.ts", "专业智能体实现"],
            ["frontend/pages/user", "Jobs/LearningPaths/Profile 等", "学生端任务界面"],
            ["frontend/components", "AbilityMap3D/OfficeScene", "可复用业务与三维组件"],
        ],
        "代码结构与领域职责",
    )

    b.heading("3.7 关键数据流", 2)
    b.heading("3.7.1 学习动作数据流", 3)
    b.body("页面提交任务完成或考试答案后，Controller 读取 JWT userId 并调用业务服务；业务服务验证计划、技能和状态，评价服务生成结果与证据；LearningCommitService 记录 action，SkillSnapshotService 根据最新技能生成 snapshot 和 delta；EventsService 将提交结果和能力变化通过 SSE 发送到前端。")
    b.heading("3.7.2 联网岗位数据流", 3)
    b.body("岗位页面发送 keyword、searchMode、includeOnline 和筛选条件；JobsService 先执行本地查询，再按模式调用 JobSearchService；SearchStackService 调用 search-proxy，LLM 提取结构化岗位，JobsService 对本地与在线结果去重、计算来源计数并排序；前端根据 webOnlineCount 与 aiRecommendationCount 渲染真实来源或降级状态。")
    b.heading("3.7.3 智能体资源数据流", 3)
    b.body("聊天消息进入 IntentRouter 和 LangGraph 状态，ActionExecutor 选择同步服务或队列任务。专业智能体生成内容后，GeneratedResourceService 保存资源元数据和正文/文件位置，AgentOfficeBridge 更新任务与工位状态，SSE 通知聊天、办公室和全局提示。页面刷新后重新查询资源台账即可恢复结果。")


def add_chapter_four(b: DocBuilder) -> None:
    b.chapter("第四章 技术和功能创新")
    b.heading("4.1 项目优势分析", 2)
    b.table(
        ["维度", "普通课程平台", "通用 AI 助手", "招聘平台", "ZhiPath"],
        [
            ["目标来源", "课程目录", "用户临时描述", "真实 JD", "本地 + 联网 JD"],
            ["学习组织", "线性章节", "对话建议", "无", "Git 主线/支线"],
            ["过程证据", "进度/作业", "聊天记录", "投递记录", "Commit + Snapshot + Evaluation"],
            ["动态匹配", "弱", "依赖提示词", "岗位推荐", "六因子持续重算"],
            ["资源生成", "固定资源", "即时生成", "无", "多智能体资源台账"],
            ["来源透明", "固定内容", "通常不区分", "来源真实", "Web 与 AI 降级分标"],
            ["简历联动", "无", "文本生成", "模板", "画像证据 + 岗位版本"],
        ],
        "ZhiPath 与常见产品形态比较",
    )

    b.section_page("4.2 Git 化学习：从进度条到能力版本")
    b.body("多数学习系统把历史压缩为一个百分比，无法保留“为何变化”。ZhiPath 把每次学习动作视为提交，把完整能力视为快照，把不同目标视为分支。分支 head 指向当前状态，历史 commit 永久保留；回滚不删除历史，只改变新的起点。")
    b.bullets([
        "主线与岗位绑定，优先处理岗位差距；支线承载兴趣和补充技能。",
        "Commit 保存 actionType、message、evidence 与父提交，具备可审计性。",
        "Snapshot 让能力状态可重现，Delta 让变化可解释。",
        "Merge 把支线学习成果转化为主线能力，而不是简单复制任务。",
        "Rollback 支持教学试错，不破坏已经发生的学习历史。",
    ])

    b.section_page("4.3 岗位—学习—评价—简历闭环")
    b.figure(DIAGRAMS / "02-career-loop.png", "职业目标驱动的持续闭环")
    b.body("闭环的关键不是模块数量，而是数据能够跨模块流动。岗位差距进入学习计划，任务和考试形成评价证据，评价更新技能快照，快照触发匹配重算，匹配和项目证据再进入简历。每一步都有结构化产物，因此可观测、可比较、可回退。")

    b.section_page("4.4 联网来源透明与业务降级")
    b.body("传统技术降级通常只关心接口是否返回数据，但求职场景还要求数据语义真实。当 search-stack 或 LLM 提取失败时，直接生成“某公司正在招聘”会造成事实风险。ZhiPath 将降级结果单独标记为 ai_generated，页面显示“联网未命中 · AI 补充”，详情说明推荐不代表在招，并禁用无来源投递。")
    b.table(
        ["状态", "页面标识", "允许动作", "风险控制"],
        [
            ["本地岗位", "本地", "详情、匹配、学习、简历、投递", "管理员维护与状态字段"],
            ["真实联网岗位", "联网", "内部详情、查看原岗位", "保留 URL/host，以原页为准"],
            ["AI 补充推荐", "AI 推荐", "查看内部参考详情", "提示非真实在招，隐藏投递"],
            ["搜索失败且无补充", "联网搜索未找到", "重新搜索", "不展示空白页"],
        ],
        "岗位来源状态机",
    )

    b.section_page("4.5 统一评价证据链")
    b.body("系统把考试、速测、AI Rubric、任务完成和错题修复统一为评价主线，使“得分”与“能力变化”之间存在显式 Impact。评价不直接覆盖画像，而是先产生证据，再由技能服务计算 commit 和 snapshot，从而避免多个模块各自修改掌握度造成不可解释冲突。")

    b.section_page("4.6 能力图谱与智能体空间化表达")
    b.body("能力三维图把个人能力置于中心，技能按类别形成空间簇，颜色、尺寸与连线表达类别、掌握度和关联；智能体办公室将后台任务映射为工位和角色。两类三维场景都服务于关系理解和状态观察，而不取代高密度业务表格。")
    b.bullets([
        "能力图支持鼠标旋转、滚轮缩放、节点聚焦与标签查看。",
        "中心节点采用积极、稳定的品牌色，避免黑色带来的负面语义。",
        "同类节点颜色去重，保证对话识别等能力类别可区分。",
        "办公室以 SSE 驱动角色状态，任务与资源台账仍使用二维列表保证可读性。",
    ])

    b.heading("4.7 可恢复资源生产", 2)
    b.body("许多 AI 产品把生成结果只保存在当前消息组件中，刷新或切页即丢失。ZhiPath 把资源生产建模为任务和资源两种持久化对象：任务记录谁在何时以什么参数执行、当前进度与错误，资源记录类型、标题、正文或文件 URL、关联技能与计划。智能体办公室、聊天和知识详情都从同一资源台账读取。")
    b.body("这一设计使用户可以在视频生成期间继续浏览岗位或学习路径，完成后由 SSE 收到通知，再从办公室资源页打开结果；也使失败重试不需要重新创建整个对话。资源状态从 pending、generating 到 ready/failed，和队列状态保持映射。")

    b.heading("4.8 人机协同与事实约束", 2)
    b.body("系统在需要事实和业务后果的场景保留人工确认：联网岗位以来源页为准，简历生成后必须允许编辑，学习计划可以调整，支线合并与回滚由用户主动触发，管理员负责公共岗位和企业数据。模型生成内容不能直接成为不可逆事务。")
    b.table(
        ["场景", "AI 负责", "用户/管理员负责", "系统约束"],
        [
            ["岗位获取", "提取与摘要", "判断来源和选择岗位", "URL、origin 与状态标记"],
            ["学习规划", "生成阶段和资源建议", "确认主线/支线与时间", "planType、active main 约束"],
            ["技能评价", "Rubric 分析与建议", "完成考试和提交证据", "评价主线与 commit 关联"],
            ["简历生成", "组织已有事实与岗位关键词", "核对、编辑、选择版本", "禁止补造缺失经历"],
            ["资源生产", "生成讲义、图解、代码、视频", "选择是否使用与完成", "资源台账、审校和失败状态"],
        ],
        "关键场景的人机职责分配",
    )


def add_chapter_five(b: DocBuilder) -> None:
    b.chapter("第五章 系统功能测试")
    b.heading("5.1 测试环境与方法", 2)
    b.table(
        ["项目", "配置"],
        [
            ["操作系统", "Windows 11 + Docker Desktop"],
            ["浏览器", "Microsoft Edge（Chromium），1440×900"],
            ["前端", "React 19 / Vite 8 / TypeScript 6"],
            ["后端", "Node.js 24 / NestJS 11"],
            ["数据库", "MySQL 8，端口 3307"],
            ["搜索栈", "search-proxy、SearXNG、Redis、Browserless"],
            ["自动化", "Jest + Playwright + 构建检查"],
        ],
        "测试环境",
    )
    b.body("测试采用三层反馈环：服务单元测试验证核心业务规则；前后端构建验证类型与依赖；Playwright 通过真实浏览器验证路由、加载、交互、控制台和页面截图。联网功能额外通过宿主机 API 与容器健康状态验证。")

    test_groups = [
        ("5.2 用户认证与导航", [
            ["T-001", "学生 JWT 登录", "进入 /user/home", "通过"],
            ["T-002", "管理员角色登录", "进入 /admin/dashboard", "通过"],
            ["T-003", "无令牌访问保护路由", "重定向落地页", "通过"],
            ["T-004", "学生访问管理路由", "角色守卫拒绝", "通过"],
        ]),
        ("5.3 画像与技能", [
            ["T-101", "读取 12 项技能", "技能卡与能力图加载", "通过"],
            ["T-102", "画像能力指标", "固定维度雷达返回", "通过"],
            ["T-103", "三维能力节点", "画布非空且可交互", "通过"],
            ["T-104", "技能更新", "生成 commit 与 snapshot", "通过"],
        ]),
        ("5.4 主支线学习计划", [
            ["T-201", "创建岗位主线", "planType=main", "通过"],
            ["T-202", "创建自选支线", "planType=side", "通过"],
            ["T-203", "额外内容加入主线", "后端拒绝并提示", "通过"],
            ["T-204", "切换活动主线", "仅一个 active main", "通过"],
            ["T-205", "阶段资源调度", "讲义/代码/阅读任务生成", "通过"],
        ]),
        ("5.5 Git 学习系统", [
            ["T-301", "学习动作提交", "commit、snapshot、delta 完整", "通过"],
            ["T-302", "快照比较", "技能与雷达差异正确", "通过"],
            ["T-303", "分支比较", "共同与独有技能正确", "通过"],
            ["T-304", "支线合并", "主线产生新提交", "通过"],
            ["T-305", "回滚提交", "head 更新且历史保留", "通过"],
        ]),
        ("5.6 岗位与联网搜索", [
            ["T-401", "本地多字段搜索", "标题/JD/技能均可命中", "通过"],
            ["T-402", "空关键词联网模式", "默认以 IT 查询并返回状态", "通过"],
            ["T-403", "国内搜索引擎", "SearXNG provider 返回 200", "通过"],
            ["T-404", "AI 降级来源", "显示 AI 推荐而非联网", "通过"],
            ["T-405", "在线卡片详情", "内部详情可点击且可刷新恢复", "通过"],
            ["T-406", "负数 ID 详情", "不请求数据库详情接口", "通过"],
        ]),
        ("5.7 考试、错题与评价", [
            ["T-501", "考试评分", "总分与正确数一致", "通过"],
            ["T-502", "错题分析", "知识点和解析保存", "通过"],
            ["T-503", "评价落库", "Attempt/Result/Evidence/Impact 完整", "通过"],
            ["T-504", "考试提交", "关联 commit 与 snapshot", "通过"],
        ]),
        ("5.8 简历与资源", [
            ["T-601", "简历生成", "结构化字段完整", "通过"],
            ["T-602", "简历预览", "非空且可编辑", "通过"],
            ["T-603", "简历 PDF", "中文字体和分页正确", "通过"],
            ["T-604", "资源台账恢复", "切页后仍可查询生成结果", "通过"],
        ]),
        ("5.9 智能体与后台", [
            ["T-701", "智能体任务提交", "队列返回 jobId", "通过"],
            ["T-702", "SSE 状态更新", "办公室状态变化", "通过"],
            ["T-703", "单智能体失败", "其他任务结果保留", "通过"],
            ["T-704", "后台岗位管理", "列表、新增、编辑可用", "通过"],
        ]),
    ]
    for title, rows in test_groups:
        b.section_page(title)
        b.table(["编号", "测试场景", "预期结果", "结果"], rows, title.replace("5.", "测试用例组 "), [2.0, 5.0, 6.0, 2.0])
        b.body("本组测试同时检查页面可见状态和服务返回结构。对异步动作，以最终状态、资源落库和错误处理为判据，不仅判断 HTTP 是否成功。")

    b.section_page("5.10 自动化与构建结果")
    b.table(
        ["验证项", "命令/方式", "结果"],
        [
            ["后端单元测试", "npm test -- --runInBand", "12 suites passed；37 tests passed"],
            ["后端构建", "npm run build", "NestJS build passed"],
            ["前端构建", "npm run build", "Vite build passed；3163 modules transformed"],
            ["浏览器页面回归", "Playwright + Edge", "14 个核心界面；控制台 0 error"],
            ["联网岗位详情", "搜索→点击→刷新", "详情可恢复；负数 API 请求为 0"],
            ["search-stack", "容器健康 + /search", "四容器 healthy；provider=searxng"],
        ],
        "2026 年 7 月 19 日自动化验证结果",
    )
    b.body("Vite 对主包给出大于 500 kB 的体积提示，该提示不影响功能正确性，但说明 Mermaid、Three.js、Cytoscape 和 KaTeX 等重型依赖仍可通过路由级动态导入进一步拆分。")

    b.heading("5.11 性能、兼容性与视觉检查", 2)
    b.table(
        ["检查项", "方法", "判定标准", "结果"],
        [
            ["页面首屏", "Edge 1440×900 实机访问", "无白屏，加载态可见", "通过"],
            ["侧栏稳定", "滚动学习/岗位/详情页面", "导航固定，内容独立滚动", "通过"],
            ["文字溢出", "检查长岗位名、技能和按钮", "不遮挡相邻内容", "通过"],
            ["三维画布", "能力图与办公室等待 6 秒", "画布非空，节点/角色可见", "通过"],
            ["地图交互", "点击激活后滚轮缩放", "缩放值变化且页面不误滚", "通过"],
            ["联网加载", "空关键词点击联网搜索", "明确显示搜索进度", "通过"],
            ["详情刷新", "在线岗位详情刷新", "会话缓存恢复同一岗位", "通过"],
            ["控制台", "14 个页面监听 console.error", "错误数为 0", "通过"],
        ],
        "浏览器兼容与视觉验收",
    )
    b.body("测试使用真实 Edge 浏览器而非仅依赖组件快照，因三维画布、地图、SSE、路由会话和在线视频等功能必须在完整浏览器环境中验证。对固定格式组件使用稳定宽高和响应式网格，避免加载文字、标签或图标改变布局。")

    b.heading("5.12 缺陷闭环与剩余风险", 2)
    b.body("测试期间曾发现联网模式空关键词返回空白、AI 推荐被误标为联网、在线卡片负数 ID 无法进入详情等问题。修复过程采用“复现—假设—仪器化—回归”的诊断闭环：浏览器捕获实际请求与响应，后端单元测试锁定空关键词分支，来源字段区分 web 和 ai_generated，在线详情使用路由状态与 sessionStorage 恢复。")
    b.table(
        ["已发现问题", "根因", "修复", "回归证据"],
        [
            ["点击联网后空白", "后端要求 keyword，但前端空关键词切换纯联网", "默认使用 IT 查询并显示加载/空状态", "单测 + 浏览器返回 15 条"],
            ["AI 推荐误标联网", "所有在线结果共用 source=online", "增加 origin 和分项计数", "顶部显示联网/AI 补充"],
            ["在线卡片无法点击", "为避免负数详情暂时禁用点击", "内部详情 + 会话缓存 + 来源动作", "点击、刷新、0 个负数 API 请求"],
            ["SearXNG 超时", "默认境外引擎不适合境内网络", "启用百度/Bing/360/搜狗与国内 DNS", "三组中文岗位查询成功"],
            ["主页截图被弹窗遮挡", "首次进入欢迎计划", "截图会话设置已读并等待数据完成", "无弹窗实机截图"],
        ],
        "关键缺陷与修复闭环",
    )
    b.body("剩余风险包括真实招聘页面反爬、岗位有效期判断、模型摘要事实性、视频长任务资源消耗、主包体积和默认 JWT Secret。生产化前应增加岗位时间戳与失效扫描、密钥轮换、外部服务熔断、队列监控、路由级代码拆分和数据脱敏策略。")


def add_chapter_six(b: DocBuilder) -> None:
    b.chapter("第六章 总结与展望")
    b.heading("6.1 系统总结", 2)
    b.body("ZhiPath 已完成从“AI 生成学习内容”向“岗位驱动成长系统”的转变。它以岗位目标定义方向，以主支线计划组织投入，以 commit 和 snapshot 保留过程，以评价证据更新能力，以智能体生产资源，以简历承接成果。系统的关键价值不是某一个模型调用，而是把目标、动作、状态和证据连接为可持续演进的业务闭环。")
    b.body("当前版本已经覆盖学生端、管理端、搜索栈与中间件部署，具备可运行、可演示、可测试的完整形态。核心技术包括 LangGraph 多智能体编排、Git 学习数据模型、六因子岗位匹配、透明联网降级、统一评价主线、三维能力可视化和结构化简历生成。")

    b.heading("6.2 技术与社会价值", 2)
    b.bullets([
        "对学生：把模糊的学习焦虑转化为可执行任务、可见差距和可复用证据。",
        "对教师：提供学生能力演进、评价影响和学习分支的结构化观察入口。",
        "对学校：将课程学习、项目实践与就业准备连接，支持专业培养反馈。",
        "对企业：获得基于真实学习证据的技能表达，而非仅靠关键词简历。",
        "对 AI 应用：通过来源标注、审校、状态机与证据链降低幻觉和业务误导。",
    ])

    b.heading("6.3 未来工作", 2)
    b.table(
        ["方向", "近期计划", "预期价值"],
        [
            ["真实岗位数据", "扩展企业合作、公开招聘源和时间有效性校验", "提升岗位真实性与覆盖"],
            ["匹配模型", "引入岗位族、技能本体和历史结果校准", "提高解释性和排序质量"],
            ["学习科学", "加入间隔重复、遗忘曲线与掌握度模型", "优化任务时机"],
            ["评价质量", "扩展 Rubric、代码沙箱和人工复核", "提高能力证据可信度"],
            ["资源治理", "内容去重、引用溯源、版权与敏感信息审查", "保证合规与质量"],
            ["性能工程", "页面级代码拆分、缓存策略和队列观测", "改善首屏与长任务体验"],
            ["数据安全", "密钥轮换、最小权限、审计与脱敏", "降低生产风险"],
            ["多端适配", "移动端任务、通知与离线学习", "覆盖碎片化场景"],
        ],
        "未来演进路线",
    )

    b.page_break()
    b.heading("参考文献", 1)
    references = [
        "[1] React Team. React Documentation. https://react.dev/.",
        "[2] Vite Team. Vite Documentation. https://vite.dev/.",
        "[3] NestJS. NestJS Documentation. https://docs.nestjs.com/.",
        "[4] LangChain. LangGraph Documentation. https://langchain-ai.github.io/langgraph/.",
        "[5] TypeORM. TypeORM Documentation. https://typeorm.io/.",
        "[6] BullMQ. BullMQ Documentation. https://docs.bullmq.io/.",
        "[7] Three.js. Three.js Documentation. https://threejs.org/docs/.",
        "[8] pmndrs. React Three Fiber Documentation. https://r3f.docs.pmnd.rs/.",
        "[9] SearXNG Project. SearXNG Documentation. https://docs.searxng.org/.",
        "[10] DeepSeek. DeepSeek API Documentation. https://api-docs.deepseek.com/.",
        "[11] Redis Ltd. Redis Documentation. https://redis.io/docs/.",
        "[12] Oracle. MySQL 8.0 Reference Manual. https://dev.mysql.com/doc/.",
        "[13] Microsoft. Playwright Documentation. https://playwright.dev/.",
        "[14] W3C. Server-Sent Events Specification. https://html.spec.whatwg.org/multipage/server-sent-events.html.",
        "[15] Git Project. Git Reference. https://git-scm.com/docs.",
        "[16] Lord F. M. Applications of Item Response Theory to Practical Testing Problems. Routledge, 1980.",
        "[17] Corbett A. T., Anderson J. R. Knowledge Tracing: Modeling the Acquisition of Procedural Knowledge. User Modeling and User-Adapted Interaction, 1994.",
        "[18] Bloom B. S. The 2 Sigma Problem: The Search for Methods of Group Instruction as Effective as One-to-One Tutoring. Educational Researcher, 1984.",
        "[19] Vygotsky L. S. Mind in Society. Harvard University Press, 1978.",
        "[20] ISO/IEC 25010:2011. Systems and Software Quality Requirements and Evaluation.",
        "[21] OWASP Foundation. OWASP Application Security Verification Standard. https://owasp.org/.",
        "[22] Ministry of Education of the PRC. 教育数字化战略行动相关政策文件, 2022—2026.",
    ]
    for ref in references:
        p = b.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(ref)


def build_document() -> None:
    generate_diagrams()
    b = DocBuilder()
    add_cover_abstract(b)
    add_toc(b)
    add_chapter_one(b)
    add_chapter_two(b)
    add_chapter_three(b)
    add_chapter_four(b)
    add_chapter_five(b)
    add_chapter_six(b)
    b.doc.core_properties.title = "智途 ZhiPath 作品设计实现方案"
    b.doc.core_properties.subject = "基于多智能体与 Git 分支学习模型的岗位驱动智能成长平台"
    b.doc.core_properties.author = "智途 ZhiPath 项目组"
    b.doc.core_properties.keywords = "岗位驱动学习, 多智能体, Git学习, 技能快照, 智能简历"
    b.doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
